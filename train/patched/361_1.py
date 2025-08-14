import docx
from lxml import etree

malicious_doc = docx.Document()
malicious_doc.add_paragraph(
    """<?xml version="1.0"?>
    <!DOCTYPE root [
        <!ENTITY xxe SYSTEM "file:///etc/passwd">
    ]>
    <root>
        <data>&xxe;</data>
    </root>"""
)
malicious_doc.save('malicious.docx')

safe_doc = docx.Document()
safe_doc.add_paragraph("This is a safe document.")
safe_doc.save('safe.docx')

doc = docx.Document('safe.docx')

xml_content = doc.part.element.xml

def safe_parse(xml_content):
    parser = etree.XMLParser(resolve_entities=False)
    return etree.fromstring(xml_content, parser)

try:
    safe_xml = safe_parse(xml_content)
    content = [p.text for p in doc.paragraphs]
    print(content)
except Exception as e:
    print(f"An error occurred: {e}")